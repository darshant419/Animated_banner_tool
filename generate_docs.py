#!/usr/bin/env python3
"""
Generate architecture documentation in Word and PDF formats
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image as RLImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
import json

def add_heading(doc, text, level=1):
    """Add a heading to the document"""
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph to the document"""
    p = doc.add_paragraph(text)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    return p

def shade_cell(cell, color):
    """Add shading to a cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_word_doc():
    """Create a Word document with architecture diagrams"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Banner Tool - Architecture Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Comprehensive System Design & Component Overview')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(12)
    subtitle_format.font.italic = True
    
    doc.add_paragraph('Generated: 2026-09-01')
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. System Overview',
        '2. High-Level Architecture',
        '3. Component Hierarchy',
        '4. State Management Structure',
        '5. Animation Pipeline',
        '6. Element Types & Rendering',
        '7. Technology Stack',
        '8. Data Flow Architecture',
        '9. Key Features & Implementations',
        '10. Design Patterns',
        '11. File Structure Organization'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. System Overview
    doc.add_heading('1. System Overview', 1)
    doc.add_paragraph(
        'The Banner Tool is a professional React-based animated banner designer built for creating '
        'compliant pharmaceutical and healthcare marketing materials. It provides a comprehensive '
        'visual design interface with advanced animation capabilities, multi-artboard support, and '
        'specialized compliance features like ISI (Important Safety Information) components.'
    )
    
    doc.add_heading('Key Capabilities', 2)
    capabilities = [
        'Multi-artboard design for multiple banner sizes simultaneously',
        'Advanced keyframe-based animation system with 300+ presets',
        'ISI (Important Safety Information) scroll components',
        'Full undo/redo history with state snapshots',
        'Real-time preview with GSAP timeline control',
        'Multiple element types: text, shapes, images, videos, HTML',
        'Export and packaging functionality',
        'Template-based workflows for EMR and animated modes'
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. High-Level Architecture
    doc.add_heading('2. High-Level Architecture', 1)
    
    doc.add_heading('Architecture Layers', 2)
    
    # Create architecture table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Components'
    hdr_cells[2].text = 'Responsibility'
    
    layers = [
        ('Entry Point', 'main.tsx, App.tsx, ModeSelect', 'Application initialization and mode selection'),
        ('State Management', 'Zustand Store, designStore.ts', 'Global state, history, and store mutations'),
        ('UI Components', 'MainLayout, Panels, Toolbar', 'User interface and interaction handling'),
        ('Canvas System', 'Konva, DesignCanvas, Renderer', 'Element rendering and visual display'),
        ('Animation Engine', 'GSAP, Keyframes, AnimationHelpers', 'Timeline control and animation playback'),
        ('Data Models', 'DesignElement, Artboard, AnimationKeyframe', 'Core data structures'),
        ('Utilities', 'keyframes.ts, animations.ts, templates.ts', 'Helper functions and utilities'),
        ('Build & Tools', 'Vite, TypeScript, Vitest, ESLint', 'Development and build infrastructure'),
    ]
    
    for layer_name, components, responsibility in layers:
        row_cells = table.add_row().cells
        row_cells[0].text = layer_name
        row_cells[1].text = components
        row_cells[2].text = responsibility
    
    doc.add_page_break()
    
    # 3. Component Hierarchy
    doc.add_heading('3. Component Hierarchy', 1)
    
    doc.add_paragraph(
        'The component structure follows a hierarchical composition pattern with MainLayout '
        'as the central orchestrator coordinating multiple specialized panels:'
    )
    doc.add_paragraph()
    
    hierarchy_text = """
App
├── ModeSelect (mode selection screen)
└── MainLayout (Central Orchestrator)
    ├── Toolbar (Tool Selection)
    │   └── text, shapes, upload, isiScroll, assets, templates, layers, variations
    ├── DesignCanvas (Konva Stage - Main Drawing Area)
    │   ├── Multiple Artboards (multi-view editing)
    │   ├── ElementShape (per element rendering)
    │   │   ├── Text Elements
    │   │   ├── Rectangle Elements
    │   │   ├── Circle Elements
    │   │   ├── Image Elements
    │   │   ├── Shape Elements (SVG Paths)
    │   │   ├── HTML Elements (iframes)
    │   │   └── Video Elements
    │   ├── ISIScroll Component (scrollable ISI panel)
    │   ├── ISIOverlay Component (ISI header bar)
    │   ├── Transformer (selection/manipulation handles)
    │   └── AnimationHelpers (GSAP playback integration)
    ├── PropertiesPanel (Edit selected element)
    │   ├── Element Properties (position, size, styling)
    │   ├── Animation Properties (keyframes, easing)
    │   ├── Canvas Properties (background, dimensions)
    │   └── Artboard Management
    ├── Timeline (Animation timeline editor)
    │   ├── Playhead Control
    │   ├── Element Animation Tracks
    │   ├── Keyframe Visual Editor
    │   └── Zoom/Pan Controls
    ├── LayersPanel (Layer management)
    │   ├── Layer List
    │   ├── Visibility Toggles
    │   └── Lock Controls
    ├── TemplatesPanel (Template library)
    │   └── Pre-built Banner Templates
    ├── AssetsPanel (Media management)
    │   └── Image/Media Asset Browser
    └── VariationsPanel (Size variations)
        └── Artboard Size Management
    """
    
    # Add as formatted code
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.25)
    for line in hierarchy_text.strip().split('\n'):
        p = doc.add_paragraph(line, style='No Spacing')
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    
    doc.add_page_break()
    
    # 4. State Management Structure
    doc.add_heading('4. State Management Structure', 1)
    doc.add_paragraph(
        'The application uses Zustand for state management with a comprehensive store '
        'that handles all application state including elements, animations, artboards, and history:'
    )
    
    state_sections = [
        ('Elements Management', [
            'elements: DesignElement[] - All design elements',
            'selectedId: string | null - Currently selected element',
            'addElement() - Add new element',
            'updateElement() - Modify element properties',
            'removeElement() - Delete element',
            'reorderElement() - Change stacking order'
        ]),
        ('Animation System', [
            'playheadTime: number - Current timeline position',
            'isPlaying: boolean - Playback state',
            'totalDuration: number - Timeline duration',
            'loop: boolean - Loop animation',
            'selectedKeyframe: SelectedKeyframe | null - Selected keyframe',
            'addKeyframe() - Add keyframe',
            'updateKeyframe() - Modify keyframe',
            'removeKeyframe() - Delete keyframe'
        ]),
        ('Artboard Management', [
            'artboards: Artboard[] - All artboards',
            'activeArtboardId: string - Currently active artboard',
            'multiArtboardView: boolean - Multi-view mode',
            'addArtboard() - Create artboard',
            'removeArtboard() - Delete artboard',
            'setActiveArtboard() - Switch active artboard'
        ]),
        ('Canvas Properties', [
            'canvasWidth: number - Canvas dimensions',
            'canvasHeight: number',
            'canvasBackground: string - Background color',
            'canvasBackgroundImage?: string - Background image',
            'setCanvasBackground() - Update background',
            'setCanvasBackgroundImage() - Update image'
        ]),
        ('History & Undo/Redo', [
            'past: HistoryState[] - Undo stack',
            'future: HistoryState[] - Redo stack',
            'undo() - Undo action',
            'redo() - Redo action',
            'clearHistory() - Clear undo/redo stacks'
        ]),
    ]
    
    for section_name, items in state_sections:
        doc.add_heading(section_name, 3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 5. Animation Pipeline
    doc.add_heading('5. Animation Pipeline', 1)
    doc.add_paragraph(
        'The animation system uses a multi-layer pipeline that converts animation presets '
        'into keyframes, which are then played back through GSAP with real-time canvas updates:'
    )
    
    doc.add_heading('Animation Flow', 2)
    animation_flow = [
        '1. User defines animation via Timeline or Properties panel',
        '2. Animation preset or keyframes stored in DesignElement',
        '3. Timeline component triggers animation playback',
        '4. AnimationHelpers.buildMasterTimeline() compiles all keyframes',
        '5. GSAP creates timeline with easing and interpolation',
        '6. Each frame update triggers Konva canvas re-render',
        '7. Canvas displays animated elements in real-time',
        '8. User can scrub playhead to preview any frame'
    ]
    for step in animation_flow:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('Animation Types Supported', 2)
    anim_types = [
        ('Keyframe Animations', 'Manual keyframe-based control with precise timing'),
        ('Preset Animations', 'Entrance, exit, and attention animations from Animista catalog'),
        ('Entrance/Exit Animations', 'Automatic animations at start/end of timeline'),
        ('Timed Animation Blocks', 'Multiple animations on one element at different times'),
        ('Hover Effects', 'Interactive animations triggered on mouseover'),
        ('Easing Functions', '13+ easing options (linear, power, bounce, elastic, etc.)')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Animation Type'
    hdr_cells[1].text = 'Description'
    
    for anim_type, description in anim_types:
        row_cells = table.add_row().cells
        row_cells[0].text = anim_type
        row_cells[1].text = description
    
    doc.add_page_break()
    
    # 6. Element Types
    doc.add_heading('6. Element Types & Rendering', 1)
    doc.add_paragraph(
        'The design canvas supports multiple element types, each with specialized properties '
        'and rendering behavior through Konva:'
    )
    
    element_types = [
        ('Text', 'Typography elements with font control, text alignment, letter spacing, line height, shadows, and full animation support'),
        ('Rect', 'Rectangle shapes with corner radius, stroke, fill colors, opacity, rotation, scale, and shadows'),
        ('Circle', 'Circular shapes with radius control, styling, and animation capabilities'),
        ('Image', 'Raster images from URLs with transform and animation support'),
        ('Video', 'Video elements with similar properties to images'),
        ('Shape', 'SVG Path-based shapes with fill/stroke and animation'),
        ('ISI Scroll', 'Special component for scrollable healthcare compliance content with auto-scroll and hover pause'),
        ('HTML', 'Iframe wrapper for custom HTML content with positioning and opacity control')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Element Type'
    hdr_cells[1].text = 'Description'
    
    for elem_type, description in element_types:
        row_cells = table.add_row().cells
        row_cells[0].text = elem_type
        row_cells[1].text = description
    
    doc.add_page_break()
    
    # 7. Technology Stack
    doc.add_heading('7. Technology Stack', 1)
    doc.add_paragraph('The application is built on modern web technologies:')
    
    tech_data = [
        ('UI Framework', 'React 19.2', 'Component-based user interface'),
        ('State Management', 'Zustand 5.0', 'Lightweight global state with subscriptions'),
        ('Canvas Rendering', 'Konva 10.0 + react-konva', '2D canvas drawing and transformation'),
        ('Animation Engine', 'GSAP 3.14', 'Professional animation timelines and playback'),
        ('Styling', 'Tailwind CSS 3.4', 'Utility-first CSS framework'),
        ('Build Tool', 'Vite 7.2', 'Fast development server and optimized builds'),
        ('Language', 'TypeScript 5.9', 'Type-safe JavaScript development'),
        ('Testing', 'Vitest 4.1', 'Unit and integration testing'),
        ('Linting', 'ESLint 9.3', 'Code quality and style checks'),
        ('Icons', 'Lucide React', 'Beautiful, consistent icon library'),
        ('Utilities', 'jszip, clsx, tailwind-merge', 'Helper libraries for specific tasks'),
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Purpose'
    
    for category, tech, purpose in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = tech
        row_cells[2].text = purpose
    
    doc.add_page_break()
    
    # 8. Data Flow
    doc.add_heading('8. Data Flow Architecture', 1)
    doc.add_paragraph('The application follows a unidirectional data flow pattern:')
    
    flow_steps = [
        ('User Interaction', 'User interacts with UI components (toolbar, properties panel, timeline, etc.)'),
        ('Event Handler', 'Component event handler captures the interaction'),
        ('Store Mutation', 'Handler calls Zustand store method to update state'),
        ('State Update', 'Zustand updates global state and notifies subscribers'),
        ('Component Re-render', 'Subscribed components receive new state and re-render'),
        ('Canvas Update', 'DesignCanvas component updates Konva stage with new element properties'),
        ('Visual Feedback', 'User sees changes immediately on canvas and in panels'),
        ('History', 'State changes are captured in undo/redo history stack')
    ]
    
    for step_name, description in flow_steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(step_name + ': ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # 9. Key Features
    doc.add_heading('9. Key Features & Implementations', 1)
    
    features = [
        ('Multi-Artboard Support', [
            'Support for multiple banner sizes (300x250, 728x90, 320x50, etc.)',
            'Each artboard maintains independent element layouts',
            'Multi-view editing mode to design multiple sizes simultaneously',
            'Pre-configured campaign size templates',
            'Easy switching between artboards'
        ]),
        ('Advanced Animation System', [
            'Keyframe-based timeline with frame-accurate control',
            '300+ preset animations from Animista.net catalog',
            'Entrance and exit animation patterns',
            'Timed animation blocks for complex sequences',
            'Interactive hover effects (color change, glow, scale, shadow)',
            'Full easing control with 13+ preset easing curves',
            'GSAP integration for smooth playback'
        ]),
        ('ISI Components (Healthcare Compliance)', [
            'Dedicated ISI Scroll component with scrollable content',
            'ISI header bar with prescribing information',
            'Auto-scroll functionality with hover pause',
            'Customizable styling (colors, fonts, padding, margins)',
            'Regulatory-compliant layout'
        ]),
        ('Export & Packaging', [
            'jszip integration for bundling assets',
            'HTML generation for standard sizes (300x250 EMR format)',
            'Asset collection and optimization',
            'Ready-for-deployment package generation'
        ]),
        ('State Persistence & History', [
            'Complete undo/redo with full state snapshots',
            'Efficient history stack management',
            'Design state preservation across sessions',
            'Quick reset to blank canvas'
        ]),
        ('Template System', [
            'Pre-built templates for EMR and animated modes',
            'Quick-start workflow with template selection',
            'Customizable template definitions',
            'Template switching in-project'
        ])
    ]
    
    for feature_name, details in features:
        doc.add_heading(feature_name, 2)
        for detail in details:
            doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_page_break()
    
    # 10. Design Patterns
    doc.add_heading('10. Design Patterns', 1)
    
    patterns = [
        ('Observer Pattern', 'Zustand store subscriptions drive UI updates automatically'),
        ('Component Composition', 'MainLayout composes multiple specialized panels and canvas'),
        ('Render Props', 'Konva components use render props for dynamic shape rendering'),
        ('Memoization', 'React.memo optimization for performance-critical components'),
        ('Factory Pattern', 'Element creation with default presets for each type'),
        ('Template Pattern', 'Animation presets defined as reusable animation specifications'),
        ('History Pattern', 'Undo/Redo implemented with full state snapshots'),
        ('Adapter Pattern', 'AnimationHelpers adapts Animista presets to keyframe format'),
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Pattern'
    hdr_cells[1].text = 'Usage'
    
    for pattern_name, usage in patterns:
        row_cells = table.add_row().cells
        row_cells[0].text = pattern_name
        row_cells[1].text = usage
    
    doc.add_page_break()
    
    # 11. File Structure
    doc.add_heading('11. File Structure Organization', 1)
    doc.add_paragraph('The project follows a clear directory structure organized by function:')
    
    structure_text = """
src/
├── main.tsx                    # React application entry point
├── App.tsx                     # Root component with mode selection
├── index.css                   # Global stylesheet
├── App.css                     # App-specific styles
│
├── components/                 # React components (organized by feature)
│   ├── Canvas/                 # Canvas rendering system
│   │   ├── DesignCanvas.tsx        # Main Konva stage component
│   │   ├── AnimationHelpers.ts     # GSAP animation utilities
│   │   ├── ISIScroll.tsx           # Scrollable ISI panel
│   │   └── ISIOverlay.tsx          # ISI header component
│   ├── Layout/
│   │   └── MainLayout.tsx      # Central layout orchestrator
│   ├── Toolbar/
│   │   └── Toolbar.tsx         # Tool selection and controls
│   ├── PropertiesPanel/
│   │   └── PropertiesPanel.tsx # Element property editor
│   ├── Timeline/
│   │   └── Timeline.tsx        # Animation timeline editor
│   ├── LayersPanel/
│   │   └── LayersPanel.tsx     # Layer management
│   ├── TemplatesPanel/
│   │   └── TemplatesPanel.tsx  # Template library
│   ├── AssetsPanel/
│   │   └── AssetsPanel.tsx     # Media asset browser
│   ├── VariationsPanel/
│   │   └── VariationsPanel.tsx # Artboard size variations
│   └── ModeSelect/
│       └── ModeSelect.tsx      # Mode selection screen
│
├── store/                      # State management
│   ├── designStore.ts          # Zustand store (primary state)
│   └── designStore.test.ts     # Store unit tests
│
├── utils/                      # Utility functions
│   ├── keyframes.ts            # Keyframe model and helpers
│   ├── keyframes.test.ts       # Keyframe tests
│   └── animations.ts           # Animation presets and catalog
│
├── templates/                  # Pre-built templates
│   └── emrTemplates.ts         # EMR template definitions
│
└── assets/                     # Static assets
    └── [images, fonts, etc.]
    """
    
    for line in structure_text.strip().split('\n'):
        p = doc.add_paragraph(line, style='No Spacing')
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    
    # Save document
    output_path = r'd:\2026\Banner_tool\Animated_banner_tool\Banner_Tool_Architecture.docx'
    doc.save(output_path)
    print(f"✓ Word document created: {output_path}")
    return output_path

def create_pdf_doc():
    """Create a PDF document with architecture information"""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Preformatted
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.lib import colors
    
    output_path = r'd:\2026\Banner_tool\Animated_banner_tool\Banner_Tool_Architecture.pdf'
    doc = SimpleDocTemplate(output_path, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    heading1_style = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=14,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=6,
        spaceBefore=6,
        fontName='Helvetica-Bold'
    )
    
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=colors.HexColor('#34495e'),
        spaceAfter=4,
        spaceBefore=4,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=12,
        spaceAfter=6,
        alignment=TA_JUSTIFY
    )
    
    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        spaceAfter=3,
        leading=11
    )
    
    story = []
    
    # Title
    story.append(Paragraph('Banner Tool', title_style))
    story.append(Paragraph('Architecture Documentation', heading1_style))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph('Comprehensive System Design & Component Overview', body_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Overview
    story.append(Paragraph('1. System Overview', heading1_style))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph(
        'The Banner Tool is a professional React-based animated banner designer built for creating compliant '
        'pharmaceutical and healthcare marketing materials. It provides a comprehensive visual design interface '
        'with advanced animation capabilities, multi-artboard support, and specialized compliance features.',
        body_style
    ))
    story.append(Spacer(1, 0.15*inch))
    
    story.append(Paragraph('Key Capabilities:', heading2_style))
    capabilities = [
        'Multi-artboard design for multiple banner sizes simultaneously',
        'Advanced keyframe-based animation system with 300+ presets',
        'ISI (Important Safety Information) scroll components',
        'Full undo/redo history with state snapshots',
        'Real-time preview with GSAP timeline control',
        'Multiple element types: text, shapes, images, videos, HTML',
        'Export and packaging functionality',
        'Template-based workflows for EMR and animated modes'
    ]
    for cap in capabilities:
        story.append(Paragraph('• ' + cap, bullet_style))
    
    story.append(Spacer(1, 0.2*inch))
    story.append(PageBreak())
    
    # Architecture Layers
    story.append(Paragraph('2. Architecture Layers', heading1_style))
    story.append(Spacer(1, 0.1*inch))
    
    layers_data = [
        ['Layer', 'Components', 'Responsibility'],
        ['Entry Point', 'main.tsx, App.tsx, ModeSelect', 'Application initialization'],
        ['State Management', 'Zustand Store, designStore.ts', 'Global state and mutations'],
        ['UI Components', 'MainLayout, Panels, Toolbar', 'User interface'],
        ['Canvas System', 'Konva, DesignCanvas, Renderer', 'Element rendering'],
        ['Animation Engine', 'GSAP, Keyframes, AnimationHelpers', 'Timeline control'],
        ['Data Models', 'DesignElement, Artboard, Keyframe', 'Core data structures'],
        ['Utilities', 'keyframes.ts, animations.ts', 'Helper functions'],
        ['Build Tools', 'Vite, TypeScript, Vitest', 'Dev infrastructure'],
    ]
    
    layers_table = Table(layers_data, colWidths=[1.2*inch, 2*inch, 2.6*inch])
    layers_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ecf0f1')]),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
    ]))
    story.append(layers_table)
    story.append(Spacer(1, 0.2*inch))
    story.append(PageBreak())
    
    # Component Hierarchy
    story.append(Paragraph('3. Component Hierarchy', heading1_style))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph(
        'The component structure follows a hierarchical composition pattern with MainLayout '
        'as the central orchestrator coordinating multiple specialized panels:',
        body_style
    ))
    story.append(Spacer(1, 0.15*inch))
    
    hierarchy = """App
└── MainLayout (Central Orchestrator)
    ├── Toolbar (Tool Selection)
    ├── DesignCanvas (Konva Stage)
    │   ├── Multiple Artboards
    │   ├── ElementShape (per element)
    │   ├── ISIScroll & ISIOverlay
    │   └── Transformer (selection)
    ├── PropertiesPanel (Edit properties)
    ├── Timeline (Animation editor)
    ├── LayersPanel (Layer management)
    ├── TemplatesPanel (Template library)
    ├── AssetsPanel (Media browser)
    └── VariationsPanel (Size variations)"""
    
    story.append(Preformatted(hierarchy, styles['Code']))
    story.append(Spacer(1, 0.2*inch))
    story.append(PageBreak())
    
    # State Management
    story.append(Paragraph('4. State Management (Zustand Store)', heading1_style))
    story.append(Spacer(1, 0.1*inch))
    
    state_sections = [
        ('Elements Management', ['elements array', 'selectedId tracking', 'Element CRUD operations']),
        ('Animation System', ['playheadTime', 'isPlaying state', 'Keyframe management']),
        ('Artboard Management', ['artboards array', 'activeArtboardId', 'Multi-view support']),
        ('Canvas Properties', ['canvasWidth/Height', 'Background settings', 'Visual properties']),
        ('History', ['past/future stacks', 'undo()/redo() methods', 'State snapshots']),
    ]
    
    for section_name, items in state_sections:
        story.append(Paragraph(section_name, heading2_style))
        for item in items:
            story.append(Paragraph('• ' + item, bullet_style))
        story.append(Spacer(1, 0.08*inch))
    
    story.append(Spacer(1, 0.2*inch))
    story.append(PageBreak())
    
    # Animation Pipeline
    story.append(Paragraph('5. Animation Pipeline', heading1_style))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph('Animation Flow:', heading2_style))
    
    animation_steps = [
        'User defines animation via Timeline or Properties panel',
        'Animation preset or keyframes stored in DesignElement',
        'Timeline triggers animation playback',
        'AnimationHelpers.buildMasterTimeline() compiles keyframes',
        'GSAP creates timeline with easing and interpolation',
        'Each frame update triggers Konva canvas re-render',
        'Canvas displays animated elements in real-time',
        'User can scrub playhead to preview any frame'
    ]
    
    for i, step in enumerate(animation_steps, 1):
        story.append(Paragraph(f'{i}. {step}', bullet_style))
    
    story.append(Spacer(1, 0.15*inch))
    story.append(Paragraph('Animation Types:', heading2_style))
    
    anim_types = [
        'Keyframe Animations - Manual, frame-accurate control',
        'Preset Animations - 300+ from Animista.net catalog',
        'Entrance/Exit Animations - Automatic timing',
        'Timed Animation Blocks - Complex sequences',
        'Hover Effects - Interactive animations',
        'Easing Functions - 13+ preset curves'
    ]
    
    for anim_type in anim_types:
        story.append(Paragraph('• ' + anim_type, bullet_style))
    
    story.append(Spacer(1, 0.2*inch))
    story.append(PageBreak())
    
    # Element Types
    story.append(Paragraph('6. Element Types Supported', heading1_style))
    story.append(Spacer(1, 0.1*inch))
    
    elements = [
        ('Text', 'Typography with font control, alignment, shadows, full animation'),
        ('Rect', 'Rectangles with corner radius, stroke, fill, opacity, transforms'),
        ('Circle', 'Circular shapes with radius control and animation'),
        ('Image', 'Raster images with transform and animation support'),
        ('Video', 'Video elements with properties similar to images'),
        ('Shape', 'SVG Path-based shapes with fill/stroke'),
        ('ISI Scroll', 'Healthcare compliance scrollable content with auto-scroll'),
        ('HTML', 'Iframe wrapper for custom HTML content'),
    ]
    
    for elem_type, desc in elements:
        story.append(Paragraph(f'<b>{elem_type}:</b> {desc}', body_style))
    
    story.append(Spacer(1, 0.2*inch))
    story.append(PageBreak())
    
    # Technology Stack
    story.append(Paragraph('7. Technology Stack', heading1_style))
    story.append(Spacer(1, 0.1*inch))
    
    tech_data = [
        ['Category', 'Technology', 'Purpose'],
        ['UI Framework', 'React 19.2', 'Component-based interface'],
        ['State Mgmt', 'Zustand 5.0', 'Global state with subscriptions'],
        ['Canvas', 'Konva 10.0', '2D canvas drawing'],
        ['Animation', 'GSAP 3.14', 'Professional animation timelines'],
        ['Styling', 'Tailwind CSS 3.4', 'Utility-first CSS'],
        ['Build Tool', 'Vite 7.2', 'Fast dev server & builds'],
        ['Language', 'TypeScript 5.9', 'Type-safe JavaScript'],
        ['Testing', 'Vitest 4.1', 'Unit & integration tests'],
        ['Linting', 'ESLint 9.3', 'Code quality checks'],
        ['Icons', 'Lucide React', 'Icon library'],
    ]
    
    tech_table = Table(tech_data, colWidths=[1.3*inch, 1.5*inch, 3*inch])
    tech_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
    ]))
    story.append(tech_table)
    
    # Build PDF
    doc.build(story)
    print(f"✓ PDF document created: {output_path}")
    return output_path

if __name__ == '__main__':
    print("Generating architecture documentation...\n")
    word_doc = create_word_doc()
    pdf_doc = create_pdf_doc()
    print(f"\n✓ Documentation generated successfully!")
    print(f"\nFiles created:")
    print(f"  - Word: {word_doc}")
    print(f"  - PDF:  {pdf_doc}")
